"""
Módulo de Adaptação de Currículo — v3
Extrai texto do PDF, adapta com Gemini, gera PDF/DOCX de 1 página.
Suporte a inglês para vagas internacionais.
"""

import os
import re

import fitz  # PyMuPDF
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .job_analyzer import _call_gemini, _extract_json

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "data", "curriculos")

COLOR_PRIMARY = (41, 65, 122)
COLOR_ACCENT = (0, 119, 181)
COLOR_SECONDARY = (80, 80, 80)
COLOR_LIGHT = (140, 140, 140)

# Section headers by language
SECTION_HEADERS = {
    "en": {
        "objective": "Professional Objective",
        "experience": "Professional Experience",
        "education": "Education",
        "tech_skills": "Technical Skills",
        "soft_skills": "Soft Skills",
        "projects": "Projects",
        "certifications": "Certifications",
        "languages": "Languages",
    },
    "pt": {
        "objective": "Objetivo Profissional",
        "experience": "Experiência Profissional",
        "education": "Formação Acadêmica",
        "tech_skills": "Habilidades Técnicas",
        "soft_skills": "Habilidades Comportamentais",
        "projects": "Projetos",
        "certifications": "Certificações",
        "languages": "Idiomas",
    },
}


def _ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def _get_headers(lang_code):
    """Returns section headers dict based on language."""
    if lang_code and lang_code.startswith("en"):
        return SECTION_HEADERS["en"]
    return SECTION_HEADERS["pt"]


def is_international_job(job):
    """Detects if a job is international (needs English resume)."""
    source = job.get("fonte", "").lower()
    location = job.get("local", "").lower()

    international_sources = ["wellfound", "remoteok", "weworkremotely"]
    if any(s in source for s in international_sources):
        return True

    brazil_pt_keywords = [
        "brasil", "brazil", "recife", "são paulo", "rio de janeiro",
        "belo horizonte", "curitiba", "porto alegre", "salvador",
        "fortaleza", "brasília", "jaboatao",
        "portugal", "lisboa", "porto", "braga", "coimbra",
    ]
    if not any(kw in location for kw in brazil_pt_keywords):
        # Check if location seems international
        if location and location not in ["", "remote", "remoto"]:
            return True

    return False


def extract_resume_text(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"  ✗ Erro ao extrair texto do PDF: {e}")
        return ""


def _validate_adapted_data(data, candidate_name):
    if not data.get("nome"):
        data["nome"] = candidate_name
    for field in ["email", "telefone", "linkedin", "localizacao", "objetivo"]:
        if not isinstance(data.get(field), str):
            data[field] = ""
    for field in ["experiencia", "formacao", "habilidades_tecnicas",
                  "habilidades_comportamentais", "idiomas", "projetos", "certificacoes"]:
        if not isinstance(data.get(field), list):
            data[field] = []
    return data


def adapt_resume_and_analyze(resume_text, job):
    """Adapta currículo com Gemini. 1 página, inglês para vagas internacionais."""
    descricao = job.get('descricao', '')
    if not descricao or len(descricao) < 20:
        descricao = f"Vaga de {job['titulo']} na empresa {job['empresa']} em {job['local']}."

    international = is_international_job(job)
    lang_instruction = (
        "The resume MUST be generated entirely in ENGLISH. Use English section names and content."
        if international else
        "Identifique o idioma da vaga. O currículo adaptado DEVE ser gerado no MESMO IDIOMA da vaga."
    )

    prompt = f"""Você é um especialista sênior em recrutamento e IA.
Faça a análise desta vaga e adapte o currículo original do candidato para ela em UMA SÓ RESPOSTA.

REGRAS CRÍTICAS:
1. {lang_instruction}
2. NÃO invente habilidades, projetos ou experiências inexistentes no currículo original.
3. Reorganize e destaque as skills do candidato que sejam mais relevantes para os requisitos da vaga.
4. O currículo DEVE caber em EXATAMENTE 1 PÁGINA. Para isso:
   - Objetivo: máximo 2 linhas, direto e com palavras-chave da vaga.
   - Experiência: máximo 3 posições mais relevantes, com no máximo 2 bullet points cada.
   - Formação: máximo 2 itens.
   - Habilidades técnicas: máximo 8 skills (as mais relevantes primeiro).
   - Habilidades comportamentais: máximo 4.
   - Projetos: máximo 2 (apenas se muito relevantes, senão omita).
   - Certificações: máximo 3 (apenas se relevantes, senão omita).
   - Idiomas: lista compacta.
5. Mantenha TODOS os dados de contato originais intactos.
6. Priorize: Experiência > Habilidades Técnicas > Formação > Projetos > Certificações.

CURRÍCULO ORIGINAL DO CANDIDATO:
{resume_text[:3000]}

VAGA:
Título: {job.get('titulo', '')}
Empresa: {job.get('empresa', '')}
Local: {job.get('local', '')}
Descrição: {descricao[:1500]}

Retorne APENAS JSON válido sem markdown no seguinte formato:
{{
    "analise": {{
        "idioma_vaga": "{'en' if international else 'pt-BR|pt-PT|en|es'}",
        "nivel": "junior|estagio|pleno",
        "requisitos_obrigatorios": ["req1"],
        "palavras_chave": ["kw1"]
    }},
    "curriculo": {{
        "nome": "nome completo",
        "email": "email",
        "telefone": "telefone",
        "linkedin": "url linkedin",
        "localizacao": "cidade, estado/país",
        "objetivo": "objetivo adaptado (max 2 linhas)",
        "experiencia": [
            {{
                "cargo": "cargo",
                "empresa": "empresa",
                "periodo": "período",
                "descricao": ["conquista 1", "conquista 2"]
            }}
        ],
        "formacao": [
            {{
                "curso": "curso",
                "instituicao": "instituição",
                "periodo": "período"
            }}
        ],
        "habilidades_tecnicas": ["skill1", "skill2"],
        "habilidades_comportamentais": ["soft skill 1"],
        "idiomas": ["idioma 1"],
        "projetos": [
            {{
                "nome": "nome",
                "descricao": "descricao curta"
            }}
        ],
        "certificacoes": ["certificação 1"]
    }}
}}"""

    print(f"  Adaptando currículo para: {job.get('titulo', '')} ({job.get('empresa', '')})...")
    if international:
        print("  🌍 Vaga internacional detectada — currículo em INGLÊS")
    response_text = _call_gemini(prompt)

    if response_text == "__RATE_LIMIT__":
        return {"_rate_limit_fallback": True}, {}

    if not response_text:
        print("  ⚠ Gemini não respondeu. Usando currículo base.")
        return {}, {}

    result = _extract_json(response_text)
    if not result:
        print("  ⚠ Não foi possível parsear resposta do Gemini.")
        return {}, {}

    analise = result.get("analise", {})
    curriculo = result.get("curriculo", {})

    # Force English for international jobs
    if international:
        analise["idioma_vaga"] = "en"

    candidate_name = os.getenv("CANDIDATE_NAME", "Paulo Antonio do Nascimento Neto")
    return _validate_adapted_data(curriculo, candidate_name), analise


# ═══════════════════════════════════════════════════════════════════════
#  GERADOR DE PDF (1 página, compacto)
# ═══════════════════════════════════════════════════════════════════════

_UNICODE_REPLACEMENTS = {
    "\u2013": "-", "\u2014": "-", "\u2018": "'", "\u2019": "'",
    "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u00b7": "-",
    "\u2022": "-", "\u2010": "-", "\u2011": "-", "\u00a0": " ",
    "\u200b": "", "\u00ad": "",
}


def _sanitize_text(text):
    if not text:
        return ""
    for unicode_char, replacement in _UNICODE_REPLACEMENTS.items():
        text = text.replace(unicode_char, replacement)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class ResumePDF(FPDF):
    """PDF de currículo compacto — otimizado para 1 página."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=False, margin=10)
        self.COLOR_PRIMARY = COLOR_PRIMARY
        self.COLOR_SECONDARY = COLOR_SECONDARY
        self.COLOR_ACCENT = COLOR_ACCENT
        self.COLOR_LIGHT = COLOR_LIGHT

    def _add_section_title(self, title):
        self.ln(2)
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*self.COLOR_PRIMARY)
        self.cell(0, 5, _sanitize_text(title.upper()), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(*self.COLOR_ACCENT)
        self.set_line_width(0.4)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(1.5)

    def _add_text(self, text, bold=False, size=8, color=None):
        style = "B" if bold else ""
        self.set_font("Helvetica", style, size)
        self.set_text_color(*(color or self.COLOR_SECONDARY))
        self.multi_cell(0, 4, _sanitize_text(text))

    def _add_bullet(self, text):
        self.set_font("Helvetica", "", 7.5)
        self.set_text_color(*self.COLOR_SECONDARY)
        bullet_x = self.l_margin + 3
        self.set_x(bullet_x)
        self.cell(3, 4, "-")
        self.set_x(bullet_x + 4)
        self.multi_cell(self.w - self.r_margin - bullet_x - 4, 4, _sanitize_text(text))

    def _fits_page(self):
        return self.get_y() < (self.h - 15)


def generate_resume_pdf(adapted_data, job, candidate_name):
    """Gera PDF de currículo de 1 página."""
    _ensure_output_dir()

    empresa_slug = re.sub(r"[^\w]", "_", job.get("empresa", "empresa"))[:30]
    vaga_slug = re.sub(r"[^\w]", "_", job.get("titulo", "vaga"))[:30]
    filename = f"CV_{candidate_name.replace(' ', '_')}_{empresa_slug}_{vaga_slug}.pdf"
    filepath = os.path.join(OUTPUT_DIR, filename)

    # Determine language for section headers
    lang = adapted_data.get("_lang", "pt")
    headers = _get_headers(lang)

    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_left_margin(12)
    pdf.set_right_margin(12)
    pdf.set_x(12)

    # ── HEADER ────────────────────────────────────────────────────
    nome = adapted_data.get("nome", candidate_name)
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(*pdf.COLOR_PRIMARY)
    pdf.cell(0, 8, _sanitize_text(nome), new_x="LMARGIN", new_y="NEXT", align="C")

    contato_parts = [p for p in [
        adapted_data.get("email", ""),
        adapted_data.get("telefone", ""),
        adapted_data.get("localizacao", ""),
    ] if p]
    if contato_parts:
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(*pdf.COLOR_LIGHT)
        pdf.cell(0, 4, _sanitize_text("  |  ".join(contato_parts)), new_x="LMARGIN", new_y="NEXT", align="C")

    if adapted_data.get("linkedin"):
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(*pdf.COLOR_ACCENT)
        pdf.cell(0, 4, _sanitize_text(adapted_data["linkedin"]), new_x="LMARGIN", new_y="NEXT", align="C")

    pdf.ln(1)
    pdf.set_draw_color(*pdf.COLOR_PRIMARY)
    pdf.set_line_width(0.6)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(1)

    # ── OBJECTIVE ─────────────────────────────────────────────────
    if adapted_data.get("objetivo") and pdf._fits_page():
        pdf._add_section_title(headers["objective"])
        pdf._add_text(adapted_data["objetivo"])

    # ── EXPERIENCE ────────────────────────────────────────────────
    exps = adapted_data.get("experiencia", [])[:3]
    if exps and pdf._fits_page():
        pdf._add_section_title(headers["experience"])
        for exp in exps:
            if not pdf._fits_page():
                break
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.set_text_color(*pdf.COLOR_SECONDARY)
            pdf.cell(0, 5, _sanitize_text(exp.get("cargo", "")), new_x="LMARGIN", new_y="NEXT")
            info = [p for p in [exp.get("empresa"), exp.get("periodo")] if p]
            if info:
                pdf.set_font("Helvetica", "I", 7.5)
                pdf.set_text_color(*pdf.COLOR_LIGHT)
                pdf.cell(0, 4, _sanitize_text(" | ".join(info)), new_x="LMARGIN", new_y="NEXT")
            for item in exp.get("descricao", [])[:2]:
                pdf._add_bullet(item)
            pdf.ln(1)

    # ── EDUCATION ─────────────────────────────────────────────────
    edus = adapted_data.get("formacao", [])[:2]
    if edus and pdf._fits_page():
        pdf._add_section_title(headers["education"])
        for edu in edus:
            if not pdf._fits_page():
                break
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.set_text_color(*pdf.COLOR_SECONDARY)
            pdf.cell(0, 5, _sanitize_text(edu.get("curso", "")), new_x="LMARGIN", new_y="NEXT")
            info = [p for p in [edu.get("instituicao"), edu.get("periodo")] if p]
            if info:
                pdf.set_font("Helvetica", "I", 7.5)
                pdf.set_text_color(*pdf.COLOR_LIGHT)
                pdf.cell(0, 4, _sanitize_text(" | ".join(info)), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

    # ── TECH SKILLS ───────────────────────────────────────────────
    skills = adapted_data.get("habilidades_tecnicas", [])[:8]
    if skills and pdf._fits_page():
        pdf._add_section_title(headers["tech_skills"])
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(*pdf.COLOR_SECONDARY)
        pdf.multi_cell(0, 4, _sanitize_text("  •  ".join(skills)))

    # ── SOFT SKILLS ───────────────────────────────────────────────
    soft = adapted_data.get("habilidades_comportamentais", [])[:4]
    if soft and pdf._fits_page():
        pdf._add_section_title(headers["soft_skills"])
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(*pdf.COLOR_SECONDARY)
        pdf.multi_cell(0, 4, _sanitize_text("  •  ".join(soft)))

    # ── PROJECTS (only if space) ──────────────────────────────────
    projs = adapted_data.get("projetos", [])[:2]
    if projs and pdf._fits_page():
        pdf._add_section_title(headers["projects"])
        for proj in projs:
            if not pdf._fits_page():
                break
            pdf.set_font("Helvetica", "B", 7.5)
            pdf.set_text_color(*pdf.COLOR_SECONDARY)
            pdf.cell(0, 4, _sanitize_text(proj.get("nome", "")), new_x="LMARGIN", new_y="NEXT")
            if proj.get("descricao"):
                pdf._add_text(proj["descricao"], size=7, color=pdf.COLOR_LIGHT)

    # ── CERTIFICATIONS (only if space) ────────────────────────────
    certs = adapted_data.get("certificacoes", [])[:3]
    if certs and pdf._fits_page():
        pdf._add_section_title(headers["certifications"])
        for cert in certs:
            if not pdf._fits_page():
                break
            pdf._add_bullet(cert)

    # ── LANGUAGES ─────────────────────────────────────────────────
    langs = adapted_data.get("idiomas", [])
    if langs and pdf._fits_page():
        pdf._add_section_title(headers["languages"])
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(*pdf.COLOR_SECONDARY)
        pdf.multi_cell(0, 4, _sanitize_text("  •  ".join(langs)))

    try:
        pdf.output(filepath)
        print(f"  ✅ PDF gerado (1 página): {filename}")
        return filepath
    except Exception as e:
        print(f"  ✗ Erro ao salvar PDF: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════════════
#  GERADOR DE DOCX (1 página, compacto)
# ═══════════════════════════════════════════════════════════════════════

def generate_resume_docx(adapted_data, job, candidate_name):
    _ensure_output_dir()

    empresa_slug = re.sub(r"[^\w]", "_", job.get("empresa", "empresa"))[:30]
    vaga_slug = re.sub(r"[^\w]", "_", job.get("titulo", "vaga"))[:30]
    filename = f"CV_{candidate_name.replace(' ', '_')}_{empresa_slug}_{vaga_slug}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    lang = adapted_data.get("_lang", "pt")
    headers = _get_headers(lang)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    def _add_heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 9)
        r, g, b = COLOR_PRIMARY if level <= 2 else COLOR_SECONDARY
        run.font.color.rgb = RGBColor(r, g, b)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        return p

    def _add_section(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*COLOR_PRIMARY)

    def _add_body(text, italic=False, size=8):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        for run in p.runs:
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(*COLOR_SECONDARY)
        return p

    _add_heading(adapted_data.get("nome", candidate_name))

    contato = [c for c in [adapted_data.get("email", ""), adapted_data.get("telefone", ""),
               adapted_data.get("localizacao", "")] if c]
    if contato:
        p = doc.add_paragraph("  |  ".join(contato))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(*COLOR_LIGHT)

    if adapted_data.get("linkedin"):
        p = doc.add_paragraph(adapted_data["linkedin"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(*COLOR_ACCENT)

    if adapted_data.get("objetivo"):
        _add_section(headers["objective"])
        _add_body(adapted_data["objetivo"])

    exps = adapted_data.get("experiencia", [])[:3]
    if exps:
        _add_section(headers["experience"])
        for exp in exps:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(exp.get("cargo", ""))
            run.bold = True
            run.font.size = Pt(9)
            info = " | ".join(filter(None, [exp.get("empresa"), exp.get("periodo")]))
            if info:
                _add_body(info, italic=True, size=7)
            for item in exp.get("descricao", [])[:2]:
                bp = doc.add_paragraph(f"• {item}")
                bp.paragraph_format.space_after = Pt(0)
                bp.paragraph_format.left_indent = Cm(0.5)
                for run in bp.runs:
                    run.font.size = Pt(7.5)

    edus = adapted_data.get("formacao", [])[:2]
    if edus:
        _add_section(headers["education"])
        for edu in edus:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(edu.get("curso", ""))
            run.bold = True
            run.font.size = Pt(9)
            info = " | ".join(filter(None, [edu.get("instituicao"), edu.get("periodo")]))
            if info:
                _add_body(info, italic=True, size=7)

    skills = adapted_data.get("habilidades_tecnicas", [])[:8]
    if skills:
        _add_section(headers["tech_skills"])
        _add_body("  •  ".join(skills), size=7.5)

    projs = adapted_data.get("projetos", [])[:2]
    if projs:
        _add_section(headers["projects"])
        for proj in projs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(proj.get("nome", ""))
            run.bold = True
            run.font.size = Pt(8)
            if proj.get("descricao"):
                _add_body(proj["descricao"], size=7)

    certs = adapted_data.get("certificacoes", [])[:3]
    if certs:
        _add_section(headers["certifications"])
        for cert in certs:
            bp = doc.add_paragraph(f"• {cert}")
            bp.paragraph_format.space_after = Pt(0)
            for run in bp.runs:
                run.font.size = Pt(7.5)

    langs = adapted_data.get("idiomas", [])
    if langs:
        _add_section(headers["languages"])
        _add_body("  •  ".join(langs), size=7.5)

    try:
        doc.save(filepath)
        print(f"  ✅ DOCX gerado (1 página): {filename}")
        return filepath
    except Exception as e:
        print(f"  ✗ Erro ao salvar DOCX: {e}")
        return ""
